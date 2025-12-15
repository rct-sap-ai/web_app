import os
import uuid
from pathlib import Path
from typing import Any
import asyncio
from jose import jwt, JWTError
from google.oauth2 import id_token
from google.auth.transport import requests as grequests
from pydantic import BaseModel
from datetime import datetime, timezone, timedelta
from fastapi import FastAPI, UploadFile, File, WebSocket, WebSocketDisconnect, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi import Header
from pathlib import Path
from docx import Document
from pypdf import PdfReader
from my_agents.paper_agents import START_AGENT

from agents import Agent, Runner
from agents.memory import OpenAIConversationsSession

from openai import OpenAI
from openai.types.responses import ResponseTextDeltaEvent
from dotenv import load_dotenv

load_dotenv()

GOOGLE_CLIENT_ID = os.getenv("GOOGLE_CLIENT_ID", "")
JWT_SECRET = os.getenv("JWT_SECRET", "")
JWT_ALG = "HS256"
ACCESS_TOKEN_MINUTES = 60

ALLOWED_EMAILS = set(
    x.strip().lower()
    for x in os.getenv("ALLOWED_EMAILS", "").split(",")
    if x.strip()
)


UPLOAD_DIR = Path(__file__).parent / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

def create_access_token(email: str) -> str:
    now = datetime.now(timezone.utc)
    payload = {
        "sub": email,
        "iat": int(now.timestamp()),
        "exp": int((now + timedelta(minutes=ACCESS_TOKEN_MINUTES)).timestamp()),
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALG)

def verify_access_token(token: str) -> str:
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALG])
        email = payload.get("sub")
        if not email:
            raise ValueError("missing sub")
        return str(email)
    except (JWTError, ValueError):
        raise HTTPException(status_code=401, detail="Invalid token")

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # lock this down later
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def route_agent(user_text: str) -> str:
    """
    Tiny agent router.
    Replace with something smarter when you add more agents.
    """
    t = user_text.lower()
    if "summarize" in t or "summary" in t:
        return "summarizer"
    if "extract" in t or "entities" in t:
        return "extractor"
    return "general"




async def call_model_streaming_agents_sdk(
    ws: WebSocket,
    agent: Agent,
    user_text: str,
    file_context: str | None,
    session: OpenAIConversationsSession,
    session_state: dict[str, Any],
):
    context_block = ""
    if file_context:
        context_block = f"\n\nUser uploaded file content:\n{file_context}\n"

    result = Runner.run_streamed(
        agent,
        input=user_text + context_block,
        session=session,
    )

    full = ""
    async for event in result.stream_events():
        if event.type == "agent_updated_stream_event":
            # update who is currently in charge
            session_state["agent"]  = event.new_agent
            await ws.send_json({"type": "status", "message": f"Agent: {event.new_agent.name}"})
            continue

        if event.type == "raw_response_event" and isinstance(event.data, ResponseTextDeltaEvent):
            delta = event.data.delta
            if delta:
                full += delta
                await ws.send_json({"type": "assistant_delta", "delta": delta})

    await ws.send_json({"type": "assistant_message", "text": full})


@app.get("/")
def root():
    return {"ok": True, "message": "Backend is running"}



def create_word_doc(pdf_path: Path) -> Path:
    if pdf_path.suffix.lower() != ".pdf":
        raise ValueError("Expected a PDF file")

    out_dir = pdf_path.parent / "generated"
    out_dir.mkdir(parents=True, exist_ok=True)

    out_path = out_dir / f"{pdf_path.stem}.docx"

    reader = PdfReader(str(pdf_path))

    doc = Document()
    doc.add_heading(f"Extracted text from {pdf_path.name}", level=1)

    any_text = False

    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = text.strip()

        doc.add_heading(f"Page {i}", level=2)

        if text:
            any_text = True
            for line in text.splitlines():
                line = line.strip()
                if line:
                    doc.add_paragraph(line)
        else:
            doc.add_paragraph("(No extractable text found on this page)")

    if not any_text:
        doc.add_paragraph("")
        doc.add_paragraph("Note: This PDF may be scanned images. Text extraction will be limited without OCR.")

    doc.save(str(out_path))
    return out_path


@app.post("/api/upload")
async def upload(file: UploadFile = File(...)) -> dict[str, Any]:
    file_id = str(uuid.uuid4())
    safe_name = (file.filename or "upload.pdf").replace("/", "_")
    dest = UPLOAD_DIR / f"{file_id}__{safe_name}"

    contents = await file.read()
    dest.write_bytes(contents)

    try:
        generated_path = create_word_doc(dest)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to create Word doc: {e}")

    return {
        "file_id": file_id,
        "filename": safe_name,
        "bytes": len(contents),
        "generated_doc": {
            "filename": generated_path.name,
            "bytes": generated_path.stat().st_size,
        },
        "download_url": f"/api/generated/{generated_path.name}",
    }



@app.get("/api/generated/{name}")
def get_generated(name: str, authorization: str | None = Header(default=None)):
    if not authorization or not authorization.lower().startswith("bearer "):
        raise HTTPException(status_code=401, detail="Missing bearer token")
    token = authorization.split(" ", 1)[1].strip()
    verify_access_token(token)

    path = (UPLOAD_DIR / "generated" / name).resolve()
    gen_dir = (UPLOAD_DIR / "generated").resolve()

    if gen_dir not in path.parents or not path.exists():
        raise HTTPException(status_code=404, detail="Not found")

    return FileResponse(str(path), filename=path.name)



@app.websocket("/ws/chat")
async def chat(ws: WebSocket):
    # your token checks...
    await ws.accept()

    session_state: dict[str, Any] = {
        "file_preview": None,
        "agent": START_AGENT,   # start here
        "session": OpenAIConversationsSession(),   # memory lives here
    }

    try:
        while True:
            msg = await ws.receive_json()
            # handle set_file etc...

            if msg.get("type") == "user_message":
                user_text = (msg.get("text") or "").strip()
                if not user_text:
                    await ws.send_json({"type": "error", "message": "Empty message."})
                    continue

                await call_model_streaming_agents_sdk(
                    ws=ws,
                    agent=session_state["agent"],
                    user_text=user_text,
                    file_context=session_state.get("file_preview"),
                    session=session_state["session"],
                    session_state=session_state,
                )
    except WebSocketDisconnect:
        return

    

class GoogleAuthBody(BaseModel):
    credential: str

@app.post("/api/auth/google")
def auth_google(body: GoogleAuthBody):
    if not GOOGLE_CLIENT_ID:
        print("Missing GOOGLE_CLIENT_ID")
        raise HTTPException(status_code=500, detail="GOOGLE_CLIENT_ID not set")
    if not JWT_SECRET:
        print("Missing JWT_SECRET")
        raise HTTPException(status_code=500, detail="JWT_SECRET not set")

    try:
        info = id_token.verify_oauth2_token(
            body.credential,
            grequests.Request(),
            GOOGLE_CLIENT_ID,
        )
    except Exception as e:
        print("Google token verify failed:", repr(e))
        raise HTTPException(status_code=401, detail="Invalid Google token")

    email = (info.get("email") or "").lower()
    email_verified = info.get("email_verified")
    print("Google login:", email, "verified:", email_verified)

    if not email or not email_verified:
        raise HTTPException(status_code=401, detail="Email not verified")

    if ALLOWED_EMAILS and email not in ALLOWED_EMAILS:
        print("Email not allowed:", email, "allowed:", ALLOWED_EMAILS)
        raise HTTPException(status_code=403, detail="Not allowed")

    token = create_access_token(email)
    return {"access_token": token, "token_type": "bearer"}
